from flask import jsonify, Flask, render_template, request, redirect, url_for, session, flash, send_file, has_request_context
import psycopg2
from psycopg2 import sql
from psycopg2.extras import RealDictCursor
from docx import Document
from io import BytesIO
import json
import os
import traceback
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from auth_utils import require_role, is_admin, is_user, generate_reset_token, send_password_reset_email, send_admin_notification
from urllib.parse import parse_qs, urlparse


app = Flask(__name__)
app.secret_key = "your_secret_key_here"

@app.route("/init-db")
def init_db():
    conn, cur = get_db()
    print("DB INIT START")
    print(f"DATABASE_URL source: environment variable present ({_database_url_source_confirmation()})")

    # 1. Create table
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        username TEXT UNIQUE,
        email TEXT UNIQUE,
        password TEXT,
        role TEXT DEFAULT 'user',
        status TEXT DEFAULT 'approved'
    );
    """)

    # 2. Create admin user
    cur.execute("""
    INSERT INTO users (username, email, password, role, status)
    VALUES ('Ch4kl3t', 'lorainenina49@gmail.com', 'l0r41n322', 'admin', 'approved')
    ON CONFLICT (username) DO UPDATE SET
        email=EXCLUDED.email,
        password=EXCLUDED.password,
        role='admin',
        status='approved';
    """)

    conn.commit()
    print("ADMIN UPSERT OK")
    cur.close()
    conn.close()

    return "DB + ADMIN CREATED"

ROLE_SESSION_KEY = 'role_sessions'
PUBLIC_ENDPOINTS = {
    'static', 'login', 'signup', 'forgot_password',
    'reset_password', 'verify_reset_token'
}
CONTEXT_AWARE_ENDPOINTS = {
    'analytics', 'api_dashboard', 'api_search_residents',
    'search_households', 'search_members', 'update_profile',
    'change_password', 'logout'
}

AUTH_SCHEMA_INITIALIZED = False

def debug_log(message):
    if os.environ.get('APP_DEBUG_LOGS', '').lower() in {'1', 'true', 'yes', 'on'}:
        print(message)



def get_db():
    url = os.environ.get("DATABASE_URL")
    if not url:
        raise Exception("DATABASE_URL not found")

    conn = psycopg2.connect(url, sslmode="require")
    return conn, conn.cursor(cursor_factory=RealDictCursor)

def _database_url_source_confirmation():
    url = os.environ.get("DATABASE_URL", "")
    parsed = urlparse(url)
    database_name = parsed.path.lstrip("/") if parsed.path else "unknown"
    return f"host={parsed.hostname or 'unknown'} db={database_name}"

def _ensure_columns(cursor, table_name, required_columns):
    cursor.execute("""
        SELECT column_name FROM information_schema.columns
        WHERE table_name=%s
    """, (table_name,))
    existing_columns = {row['column_name'] for row in cursor.fetchall()}
    for column, definition in required_columns.items():
        if column not in existing_columns:
            cursor.execute(
                sql.SQL("ALTER TABLE {} ADD COLUMN {} {}").format(
                    sql.Identifier(table_name),
                    sql.Identifier(column),
                    sql.SQL(definition)
                )
            )

def _month_filter_sql(column_name):
    return f"EXTRACT(MONTH FROM {column_name}) = %s"

def _current_year_filter_sql(column_name):
    return f"EXTRACT(YEAR FROM {column_name}) = EXTRACT(YEAR FROM CURRENT_DATE)"

def ensure_auth_schema_safe():
    """Initialize the PostgreSQL schema in one transaction.

    Canonical tables created here:
    users, password_resets, audit_logs, household.
    """
    global AUTH_SCHEMA_INITIALIZED
    if AUTH_SCHEMA_INITIALIZED:
        return True

    conn = None
    try:
        conn, cursor = get_db()
        print("DB INIT START")
        print(f"DATABASE_URL source: environment variable present ({_database_url_source_confirmation()})")

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                username VARCHAR(255) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                email VARCHAR(255),
                role VARCHAR(20) DEFAULT 'user',
                status VARCHAR(20) DEFAULT 'approved',
                signup_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        _ensure_columns(cursor, 'users', {
            'username': 'VARCHAR(255)',
            'password': 'VARCHAR(255)',
            'email': 'VARCHAR(255)',
            'role': "VARCHAR(20) DEFAULT 'user'",
            'status': "VARCHAR(20) DEFAULT 'approved'",
            'signup_date': 'TIMESTAMP'
        })

        cursor.execute("UPDATE users SET role='user' WHERE role IS NULL OR role=''")
        cursor.execute("UPDATE users SET status='approved' WHERE status IS NULL OR status=''")
        cursor.execute("""
            INSERT INTO users (username, email, password, role, status)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (username) DO UPDATE SET
                email=EXCLUDED.email,
                password=EXCLUDED.password,
                role=EXCLUDED.role,
                status=EXCLUDED.status
        """, ('Ch4kl3t', 'lorainenina49@gmail.com', 'l0r41n322', 'admin', 'approved'))
        print("ADMIN UPSERT OK")

        cursor.execute("SELECT * FROM users WHERE username=%s", ('Ch4kl3t',))
        admin_user = cursor.fetchone()
        print("ADMIN VERIFY OK" if admin_user else "ADMIN VERIFY FAILED")

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS household (
                id SERIAL PRIMARY KEY,
                firstname VARCHAR(255),
                middlename VARCHAR(255),
                surname VARCHAR(255),
                house_number VARCHAR(100),
                address TEXT,
                age VARCHAR(20),
                birthdate VARCHAR(50),
                gender VARCHAR(50),
                civil_status VARCHAR(100),
                occupation VARCHAR(255),
                household_id INTEGER,
                status VARCHAR(50) DEFAULT 'Active',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        _ensure_columns(cursor, 'household', {
            'firstname': 'VARCHAR(255)',
            'middlename': 'VARCHAR(255)',
            'surname': 'VARCHAR(255)',
            'house_number': 'VARCHAR(100)',
            'address': 'TEXT',
            'age': 'VARCHAR(20)',
            'birthdate': 'VARCHAR(50)',
            'gender': 'VARCHAR(50)',
            'civil_status': 'VARCHAR(100)',
            'occupation': 'VARCHAR(255)',
            'household_id': 'INTEGER',
            'status': "VARCHAR(50) DEFAULT 'Active'",
            'created_at': 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP'
        })

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS households (
                id SERIAL PRIMARY KEY,
                surname VARCHAR(255),
                house_number VARCHAR(100),
                address TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        _ensure_columns(cursor, 'households', {
            'surname': 'VARCHAR(255)',
            'house_number': 'VARCHAR(100)',
            'address': 'TEXT',
            'created_at': 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP'
        })

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS audit_logs (
                id SERIAL PRIMARY KEY,
                user_id VARCHAR(100),
                username VARCHAR(255),
                action_type VARCHAR(100),
                target_type VARCHAR(100) DEFAULT 'System',
                target_id VARCHAR(100) DEFAULT 'N/A',
                old_value TEXT,
                new_value TEXT,
                details TEXT,
                household_context TEXT DEFAULT 'N/A',
                status VARCHAR(30) DEFAULT 'SUCCESS',
                ip_address VARCHAR(80),
                user_agent VARCHAR(255),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        _ensure_columns(cursor, 'audit_logs', {
            'user_id': "VARCHAR(100) NULL",
            'username': 'VARCHAR(255)',
            'action_type': 'VARCHAR(100)',
            'target_type': "VARCHAR(100) DEFAULT 'System'",
            'target_id': "VARCHAR(100) DEFAULT 'N/A'",
            'old_value': 'TEXT NULL',
            'new_value': 'TEXT NULL',
            'details': 'TEXT',
            'household_context': "TEXT DEFAULT 'N/A'",
            'status': "VARCHAR(30) DEFAULT 'SUCCESS'",
            'ip_address': 'VARCHAR(80) NULL',
            'user_agent': 'VARCHAR(255) NULL',
            'created_at': 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP'
        })

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS password_resets (
                id SERIAL PRIMARY KEY,
                username VARCHAR(100) NOT NULL,
                token VARCHAR(255) UNIQUE NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                expires_at TIMESTAMP NOT NULL,
                used BOOLEAN DEFAULT FALSE,
                used_at TIMESTAMP NULL,
                ip_address VARCHAR(80),
                attempt_count INT DEFAULT 1
            )
        """)
        _ensure_columns(cursor, 'password_resets', {
            'username': 'VARCHAR(100)',
            'token': 'VARCHAR(255)',
            'created_at': 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP',
            'expires_at': 'TIMESTAMP',
            'used': 'BOOLEAN DEFAULT FALSE',
            'used_at': 'TIMESTAMP NULL',
            'ip_address': 'VARCHAR(80)',
            'attempt_count': 'INT DEFAULT 1'
        })
        
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_users_username ON users (username)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_users_status ON users (status)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_password_resets_token ON password_resets (token)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_token ON password_resets (token)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_username ON password_resets (username)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_used ON password_resets (used)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_household_id ON household (household_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_status ON household (status)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_surname ON household (surname)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_gender ON household (gender)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_created_at ON household (created_at)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_households_surname ON households (surname)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_households_house_number ON households (house_number)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_households_created_at ON households (created_at)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_created_at ON audit_logs (created_at)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_target ON audit_logs (target_type, target_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_target_action_created ON audit_logs (target_type, action_type, target_id, created_at)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_action_created ON audit_logs (action_type, created_at)")
        conn.commit()
        try:
            cursor.execute("CREATE EXTENSION IF NOT EXISTS pg_trgm")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_household_name_trgm ON household USING GIN ((CONCAT_WS(' ', firstname, middlename, surname)) gin_trgm_ops)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_households_search_trgm ON households USING GIN ((CONCAT_WS(' ', surname, house_number, address)) gin_trgm_ops)")
            conn.commit()
        except Exception as index_error:
            print(f"Warning: pg_trgm indexes skipped: {index_error}")
            conn.rollback()
        conn.close()
        AUTH_SCHEMA_INITIALIZED = True
        return True
    except Exception as e:
        print(f"Warning: Schema initialization error: {e}")
        if conn:
            conn.rollback()
            conn.close()
        return False

def ensure_auth_schema():
    """Wrapper function for backward compatibility - calls safe version."""
    return ensure_auth_schema_safe()

def role_home_endpoint():
    return 'home' if session.get('role') == 'admin' else 'user_home'

def role_profile_endpoint():
    return 'profile' if session.get('role') == 'admin' else 'user_profile'

def _store_role_identity(role, username, user_id):
    role_sessions = dict(session.get(ROLE_SESSION_KEY) or {})
    role_sessions[role] = {
        'username': username,
        'user_id': user_id,
        'role': role
    }
    session[ROLE_SESSION_KEY] = role_sessions
    session['active_role'] = role
    session['username'] = username
    session['user_id'] = user_id
    session['role'] = role
    session.modified = True

def _apply_role_identity(role):
    identity = (session.get(ROLE_SESSION_KEY) or {}).get(role)
    if not identity:
        return False

    session['username'] = identity.get('username')
    session['user_id'] = identity.get('user_id')
    session['role'] = identity.get('role', role)
    session['active_role'] = role
    return True

def _role_from_referrer():
    referrer = request.referrer or ''
    if not referrer:
        return None

    parsed = urlparse(referrer)
    referrer_role = (parse_qs(parsed.query).get('as') or [None])[0]
    if referrer_role in {'admin', 'user'}:
        return referrer_role
    if parsed.path.startswith('/user_'):
        return 'user'
    return None

def _requested_role_context():
    requested = request.args.get('as')
    if requested in {'admin', 'user'}:
        return requested

    endpoint = request.endpoint or ''
    if endpoint.startswith('user_') or request.path.startswith('/user_'):
        return 'user'

    if endpoint in CONTEXT_AWARE_ENDPOINTS:
        return _role_from_referrer() or session.get('active_role') or session.get('role')

    if endpoint and endpoint not in PUBLIC_ENDPOINTS:
        return 'admin'

    return None

@app.before_request
def restore_role_context():
    role = _requested_role_context()
    if role in {'admin', 'user'}:
        _apply_role_identity(role)

def external_reset_url(token):
    app_base_url = os.getenv('APP_BASE_URL', '').strip().rstrip('/')
    if app_base_url:
        return f"{app_base_url}{url_for('reset_password', token=token)}"
    return url_for('reset_password', token=token, _external=True)

def ensure_audit_log_schema():
    """Keep the audit log table compatible with the richer audit trail UI."""
    conn, cursor = get_db()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS audit_logs (
            id SERIAL PRIMARY KEY,
            user_id VARCHAR(100),
            username VARCHAR(255),
            action_type VARCHAR(100),
            target_type VARCHAR(100) DEFAULT 'System',
            target_id VARCHAR(100) DEFAULT 'N/A',
            old_value TEXT,
            new_value TEXT,
            details TEXT,
            household_context TEXT DEFAULT 'N/A',
            status VARCHAR(30) DEFAULT 'SUCCESS',
            ip_address VARCHAR(80),
            user_agent VARCHAR(255),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    try:
        cursor.execute("""
            SELECT column_name FROM information_schema.columns 
            WHERE table_name='audit_logs'
        """)
        existing = {row['column_name'] for row in cursor.fetchall()}
    except:
        existing = set()
    
    columns_to_add = {
        'user_id': "VARCHAR(100) NULL",
        'target_type': "VARCHAR(100) DEFAULT 'System'",
        'target_id': "VARCHAR(100) DEFAULT 'N/A'",
        'old_value': "TEXT NULL",
        'new_value': "TEXT NULL",
        'status': "VARCHAR(30) DEFAULT 'SUCCESS'",
        'ip_address': "VARCHAR(80) NULL",
        'user_agent': "VARCHAR(255) NULL"
    }
    for column, definition in columns_to_add.items():
        if column not in existing:
            cursor.execute(f"ALTER TABLE audit_logs ADD COLUMN {column} {definition}")
    conn.commit()
    conn.close()

def _audit_json(value):
    if value is None:
        return None
    if isinstance(value, str):
        return value
    return json.dumps(value, default=str, ensure_ascii=False)

def log_audit(username, action_type, details, household_context='N/A',
              user_id=None, target_type='System', target_id='N/A',
              old_value=None, new_value=None, status='SUCCESS'):
    """Log an action to the audit_logs table"""
    try:
        ensure_audit_log_schema()
        conn, cursor = get_db()
        cursor.execute("""
            INSERT INTO audit_logs (
                user_id, username, action_type, target_type, target_id,
                old_value, new_value, details, household_context,
                status, ip_address, user_agent
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            user_id or username,
            username,
            action_type,
            target_type,
            target_id,
            _audit_json(old_value),
            _audit_json(new_value),
            details,
            household_context,
            status,
            request.remote_addr if has_request_context() else None,
            (request.user_agent.string[:255] if has_request_context() and request.user_agent else None)
        ))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error logging audit: {e}")

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        try:
            # Ensure schema exists before attempting login
            ensure_auth_schema_safe()
            
            username = request.form.get('username', '').strip()
            password = request.form.get('password', '').strip()
            
            if not username or not password:
                flash("Username and password are required", "danger")
                return redirect(url_for('login'))
            
            conn, cursor = get_db()
            cursor.execute("SELECT * FROM users WHERE username=%s", (username,))
            user = cursor.fetchone()
            conn.close()
            
            if user and user['password'] == password:
                # Check account status
                status = (user['status'] or 'approved').lower()
                if status == 'pending':
                    flash("Your account is pending admin approval. Please wait for confirmation.", "warning")
                    return redirect(url_for('login'))
                elif status == 'rejected':
                    flash("Your account registration was rejected. Contact administrator for more details.", "danger")
                    return redirect(url_for('login'))
                
                role = (user['role'] or 'user').lower()
                if role not in {'admin', 'user'}:
                    role = 'user'

                # Set the full session before audit/redirect logic. The 500 error
                # can happen when code reads session["role"] before Flask has a
                # complete session for the current request, so use the verified
                # local role and safe session access below.
                session['role'] = role
                session['username'] = username
                session['user_id'] = user['id']
                _store_role_identity(role, username, user['id'])

                try:
                    log_audit(
                        username,
                        'LOGIN',
                        f'User logged in with role: {session.get("role", "user")}',
                        user_id=user['id']
                    )
                except Exception as audit_error:
                    print(f"Login audit skipped: {audit_error}")
                
                # Route based on role
                if session.get('role', 'user') == 'admin':
                    return redirect(url_for('home'))
                else:
                    return redirect(url_for('user_home'))
            else:
                flash("Invalid username or password", "danger")
                return redirect(url_for('login'))
        except Exception as e:
            print(f"Login error: {e}")
            flash(f"Login failed: {str(e)}", "danger")
            return redirect(url_for('login'))
    return render_template('login.html')

@app.route('/logout')
def logout():
    role = session.get('role')
    role_sessions = dict(session.get(ROLE_SESSION_KEY) or {})
    if role in role_sessions:
        role_sessions.pop(role, None)
        session[ROLE_SESSION_KEY] = role_sessions

    session.pop('username', None)
    session.pop('user_id', None)
    session.pop('role', None)
    session.pop('active_role', None)

    if role_sessions:
        fallback_role = 'admin' if 'admin' in role_sessions else next(iter(role_sessions))
        session[ROLE_SESSION_KEY] = role_sessions
        _apply_role_identity(fallback_role)

    return redirect(url_for('login'))

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    """Public self-registration page - no login required"""
    if request.method == 'POST':
        ensure_auth_schema()
        fullname = request.form.get('fullname', '').strip()
        email = request.form.get('email', '').strip()
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        purpose = request.form.get('purpose', '').strip()

        # Validation
        if not all([fullname, email, username, password, purpose]):
            flash('All fields are required', 'danger')
            return redirect(url_for('signup'))

        if len(username) < 3:
            flash('Username must be at least 3 characters', 'danger')
            return redirect(url_for('signup'))

        if len(password) < 8:
            flash('Password must be at least 8 characters', 'danger')
            return redirect(url_for('signup'))

        if password != confirm_password:
            flash('Passwords do not match', 'danger')
            return redirect(url_for('signup'))

        conn, cursor = get_db()
        try:
            # Check if username exists
            cursor.execute("SELECT id FROM users WHERE username=%s", (username,))
            if cursor.fetchone():
                flash('Username already exists', 'danger')
                conn.close()
                return redirect(url_for('signup'))

            # Check if email is already registered
            cursor.execute("SELECT id FROM users WHERE email=%s", (email,))
            if cursor.fetchone():
                flash('Email already registered', 'danger')
                conn.close()
                return redirect(url_for('signup'))

            # Insert new user with 'pending' status
            cursor.execute("""
                INSERT INTO users (username, password, email, role, status, signup_date) 
                VALUES (%s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
                RETURNING id
            """, (username, password, email, 'user', 'pending'))
            
            new_user_id = cursor.fetchone()['id']
            conn.commit()

            # Log the signup
            log_audit(
                'system',
                'SIGNUP',
                f'New user registration: {username} ({email}) - Status: pending',
                user_id=new_user_id,
                target_type='User',
                target_id=new_user_id,
                new_value={
                    'username': username,
                    'email': email,
                    'purpose': purpose,
                    'status': 'pending'
                }
            )

            conn.close()
            flash('Account created! Your request is pending admin review. You will receive an email notification once approved.', 'success')
            return redirect(url_for('login'))

        except Exception as e:
            conn.close()
            flash(f'Error creating account: {str(e)}', 'danger')
            return redirect(url_for('signup'))

    return render_template('signup.html')

@app.route('/api/user/status/<int:user_id>/<action>', methods=['POST'])
@require_role('admin')
def update_user_status(user_id, action):
    """Admin approve/reject pending signups"""
    conn = None
    try:
        ensure_auth_schema()
        if action not in ['approve', 'reject']:
            return jsonify({'error': 'Invalid action'}), 400

        conn, cursor = get_db()
        
        # Get user details
        cursor.execute("SELECT id, username, email, status FROM users WHERE id=%s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return jsonify({'error': 'User not found'}), 404

        old_status = user['status']
        new_status = 'approved' if action == 'approve' else 'rejected'

        # Update status
        cursor.execute(
            "UPDATE users SET status=%s WHERE id=%s",
            (new_status, user_id)
        )
        conn.commit()

        # Log the action
        log_audit(
            session.get('username', 'admin'),
            'UPDATE',
            f'User account {action}ed: {user["username"]}',
            user_id=session.get('user_id'),
            target_type='User',
            target_id=user_id,
            old_value={'status': old_status},
            new_value={'status': new_status}
        )

        return jsonify({
            'success': True,
            'message': f'User {action}ed successfully',
            'new_status': new_status
        }), 200

    except Exception as e:
        print(f"[ERROR] /api/user/status error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    """Handle password reset request"""
    if request.method == 'POST':
        ensure_auth_schema()
        email = request.form.get('email', '').strip().lower()
        
        if not email:
            flash('Please enter your email address', 'danger')
            return redirect(url_for('forgot_password'))
        
        conn, cursor = get_db()
        cursor.execute("""
            SELECT id, username, email
            FROM users
            WHERE LOWER(email)=%s AND COALESCE(status, 'approved')='approved'
        """, (email,))
        user = cursor.fetchone()
        
        if not user:
            # For security, don't reveal if user exists
            flash('If that email is registered, a reset link has been sent.', 'info')
            conn.close()
            return redirect(url_for('login'))

        username = user['username']
        
        # Generate reset token
        reset_token = generate_reset_token()
        expires_at = datetime.now() + timedelta(hours=1)
        
        try:
            cursor.execute("""
                INSERT INTO password_resets (username, token, expires_at, ip_address)
                VALUES (%s, %s, %s, %s)
            """, (username, reset_token, expires_at, request.remote_addr))
            conn.commit()
            
            # Create reset link
            reset_link = external_reset_url(reset_token)
            user_email = user['email']
            
            # Send email
            success, message = send_password_reset_email(user_email, username, reset_link)
            conn.close()
            
            if success:
                flash('If that email is registered, a reset link has been sent.', 'success')
                log_audit('System', 'PASSWORD_RESET_REQUEST', f'Password reset requested for user: {username}', 
                         target_type='User', target_id=str(user['id']))
            else:
                flash(f'Could not send email: {message}', 'warning')
            
            return redirect(url_for('login'))
        
        except Exception as e:
            conn.close()
            flash(f'Error processing reset request: {str(e)}', 'danger')
            return redirect(url_for('forgot_password'))
    
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """Handle password reset with token"""
    ensure_auth_schema()
    conn, cursor = get_db()
    
    # Verify token exists and is not expired
    cursor.execute("""
        SELECT username, expires_at, used FROM password_resets 
        WHERE token=%s AND used=FALSE
    """, (token,))
    reset_record = cursor.fetchone()
    
    if not reset_record:
        conn.close()
        flash('Invalid or expired reset link', 'danger')
        return redirect(url_for('login'))
    
    # Check if token is expired
    if datetime.fromisoformat(str(reset_record['expires_at'])) < datetime.now():
        conn.close()
        flash('Reset link has expired. Please request a new one.', 'danger')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        new_password = request.form.get('new_password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        username = reset_record['username']
        
        if not new_password or not confirm_password:
            flash('Please enter a password', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        if new_password != confirm_password:
            flash('Passwords do not match', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        if len(new_password) < 8:
            flash('Password must be at least 8 characters', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        try:
            # Update password
            cursor.execute("UPDATE users SET password=%s WHERE username=%s", (new_password, username))
            
            # Mark token as used
            cursor.execute("""
                UPDATE password_resets SET used=TRUE, used_at=CURRENT_TIMESTAMP WHERE token=%s
            """, (token,))
            
            conn.commit()
            
            # Log the reset
            cursor.execute("SELECT id FROM users WHERE username=%s", (username,))
            user = cursor.fetchone()
            log_audit('System', 'PASSWORD_RESET_SUCCESS', f'Password reset completed for: {username}', 
                     target_type='User', target_id=str(user['id']) if user else 'N/A')
            
            # Notify admin
            cursor.execute("SELECT email FROM users WHERE role='admin' LIMIT 1")
            admin = cursor.fetchone()
            if admin and admin['email']:
                send_admin_notification(admin['email'], username, 'User reset password successfully')
            
            conn.close()
            flash('Password has been reset successfully. You can now log in.', 'success')
            return redirect(url_for('login'))
        
        except Exception as e:
            conn.close()
            flash(f'Error resetting password: {str(e)}', 'danger')
            return redirect(url_for('reset_password', token=token))
    
    conn.close()
    return render_template('reset_password.html', token=token)

@app.route('/user-management')
@require_role('admin')
def user_management():
    """Admin panel for managing users"""
    ensure_auth_schema()
    conn, cursor = get_db()
    try:
        cursor.execute("SELECT id, username, role, status FROM users ORDER BY username ASC")
        users = cursor.fetchall()
        
        # Get pending signups
        cursor.execute("""
            SELECT id, username, email, signup_date FROM users 
            WHERE status='pending' 
            ORDER BY signup_date DESC
        """)
        pending_users = cursor.fetchall()
        pending_count = len(pending_users) if pending_users else 0
        
        # PostgreSQL requires selected non-aggregate columns to be grouped.
        # The old query selected raw created_at while grouping by DATE(created_at),
        # which crashed User Management in production.
        cursor.execute("""
            SELECT username, created_at, used_at, attempts FROM (
                SELECT
                    username,
                    MIN(created_at) AS created_at,
                    MAX(used_at) AS used_at,
                    COUNT(*) AS attempts
                FROM password_resets
                GROUP BY username, DATE(created_at)
            ) AS subq
            ORDER BY created_at DESC LIMIT 50
        """)
        reset_history = cursor.fetchall()
        conn.close()
        
        return render_template('user_management.html', users=users, pending_users=pending_users, 
                             pending_count=pending_count, reset_history=reset_history)
    except Exception as e:
        conn.close()
        print(f"User management error: {e}")
        flash("User Management could not load. Database schema was refreshed; please try again.", "danger")
        return redirect(url_for('home'))

@app.route('/api/pending-signups-count')
@require_role('admin')
def api_pending_signups_count():
    """API endpoint to get count of pending signups (admin only)"""
    conn = None
    try:
        ensure_auth_schema()
        conn, cursor = get_db()
        cursor.execute("SELECT COUNT(*) as count FROM users WHERE status='pending'")
        result = cursor.fetchone()
        
        return jsonify({'count': result.get('count', 0) if result else 0}), 200
    except Exception as e:
        print(f"[ERROR] /api/pending-signups-count error: {e}")
        return jsonify({'count': 0, 'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/user/role', methods=['POST'])
@require_role('admin')
def api_update_user_role():
    """API endpoint to update user role (admin only)"""
    conn = None
    try:
        ensure_auth_schema()
        data = request.get_json()
        user_id = data.get('user_id')
        new_role = data.get('role')
        
        if new_role not in {'admin', 'user'}:
            return jsonify({'error': 'Invalid role'}), 400
        
        conn, cursor = get_db()
        cursor.execute("SELECT username, role FROM users WHERE id=%s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return jsonify({'error': 'User not found'}), 404

        if str(user_id) == str(session.get('user_id')) and new_role != 'admin':
            return jsonify({'error': 'You cannot demote your own active admin account'}), 400
        
        cursor.execute("UPDATE users SET role=%s WHERE id=%s", (new_role, user_id))
        conn.commit()
        
        log_audit(session.get('username', 'admin'), 'UPDATE', f'User role changed: {user["username"]} -> {new_role}',
                 target_type='User', target_id=str(user_id), old_value=f'role: {user["role"]}', new_value=f'role: {new_role}')
        
        return jsonify({'success': True, 'message': f'User role updated to {new_role}'}), 200
    
    except Exception as e:
        print(f"[ERROR] /api/user/role error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if conn:
            conn.close()

@app.route('/register-user', methods=['GET', 'POST'])
@require_role('admin')
def register_user():
    """Admin page to register new user accounts"""
    if request.method == 'POST':
        ensure_auth_schema()
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        role = request.form.get('role', 'user').strip()

        # Validation
        if not all([username, email, password, confirm_password, role]):
            flash('All fields are required', 'danger')
            return redirect(url_for('register_user'))

        if len(username) < 3:
            flash('Username must be at least 3 characters', 'danger')
            return redirect(url_for('register_user'))

        if len(password) < 8:
            flash('Password must be at least 8 characters', 'danger')
            return redirect(url_for('register_user'))

        if password != confirm_password:
            flash('Passwords do not match', 'danger')
            return redirect(url_for('register_user'))

        if role not in ['admin', 'user']:
            flash('Invalid role selected', 'danger')
            return redirect(url_for('register_user'))

        conn, cursor = get_db()
        try:
            # Check if username exists
            cursor.execute("SELECT id FROM users WHERE username=%s", (username,))
            if cursor.fetchone():
                flash('Username already exists', 'danger')
                conn.close()
                return redirect(url_for('register_user'))

            # Insert new user
            cursor.execute("""
                INSERT INTO users (username, password, email, role) 
                VALUES (%s, %s, %s, %s)
                RETURNING id
            """, (username, password, email, role))
            
            new_user_id = cursor.fetchone()['id']
            conn.commit()

            # Log the action
            log_audit(
                session.get('username', 'admin'),
                'ADD',
                f'Created new user account: {username} (role: {role})',
                'N/A',
                user_id=session.get('user_id'),
                target_type='User',
                target_id=new_user_id,
                new_value={
                    'username': username,
                    'email': email,
                    'role': role
                }
            )

            conn.close()
            flash(f'User account "{username}" created successfully with role: {role}', 'success')
            return redirect(url_for('user_management'))

        except Exception as e:
            conn.close()
            flash(f'Error creating user: {str(e)}', 'danger')
            return redirect(url_for('register_user'))

    return render_template('register_user.html')

@app.route('/profile')
@require_role('admin')
def profile():
    if 'username' not in session:
        return redirect(url_for('login'))
    return render_template('profile.html')

@app.route('/update_profile', methods=['POST'])
@require_role('admin', 'user')
def update_profile():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    email = request.form.get('email', '').strip()
    
    if not email:
        flash("Email address is required", "danger")
        return redirect(url_for(role_profile_endpoint()))

    conn, cursor = get_db()
    cursor.execute("UPDATE users SET email=%s WHERE username=%s", (email, session.get('username')))
    conn.commit()
    conn.close()

    flash("Profile updated successfully", "success")
    return redirect(url_for(role_profile_endpoint()))

@app.route('/change_password', methods=['POST'])
@require_role('admin', 'user')
def change_password():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    current_password = request.form.get('current_password', '')
    new_password = request.form.get('new_password', '')
    confirm_password = request.form.get('confirm_password', '')
    
    if new_password != confirm_password:
        flash("New passwords do not match", "danger")
        return redirect(url_for(role_profile_endpoint()))
    
    if len(new_password) < 8:
        flash("New password must be at least 8 characters", "danger")
        return redirect(url_for(role_profile_endpoint()))

    conn, cursor = get_db()
    cursor.execute("SELECT password FROM users WHERE username=%s", (session.get('username'),))
    user = cursor.fetchone()
    if not user or user['password'] != current_password:
        conn.close()
        flash("Current password is incorrect", "danger")
        return redirect(url_for(role_profile_endpoint()))

    cursor.execute("UPDATE users SET password=%s WHERE username=%s", (new_password, session.get('username')))
    conn.commit()
    conn.close()

    flash("Password changed successfully", "success")
    return redirect(url_for(role_profile_endpoint()))

# ===== USER ROUTES (SEPARATE FROM ADMIN) =====

@app.route('/user_home')
@require_role('user')
def user_home():
    """User dashboard page - read-only community statistics"""
    if 'username' not in session:
        return redirect(url_for('login'))
    
    conn, cursor = get_db()
    debug_log("[DEBUG][statistics] Loading user dashboard totals from household and households tables")
    cursor.execute("""
        SELECT
            COUNT(*) FILTER (WHERE COALESCE(status, 'Active') != 'Deceased') AS total_residents,
            COUNT(*) FILTER (WHERE COALESCE(status, 'Active') = 'Deceased') AS total_deceased,
            COUNT(*) FILTER (WHERE gender = 'Male' AND COALESCE(status, 'Active') != 'Deceased') AS total_males,
            COUNT(*) FILTER (WHERE gender = 'Female' AND COALESCE(status, 'Active') != 'Deceased') AS total_females
        FROM household
    """)
    resident_totals = cursor.fetchone()
    total_residents = resident_totals['total_residents']
    total_deceased = resident_totals['total_deceased']
    total_males = resident_totals['total_males']
    total_females = resident_totals['total_females']
    cursor.execute("SELECT COUNT(*) AS total FROM households")
    total_households = cursor.fetchone()['total']
    debug_log(f"[DEBUG][statistics] user totals residents={total_residents} deceased={total_deceased} households={total_households} male={total_males} female={total_females}")
    conn.close()
    
    return render_template(
        'user_home.html',
        total_residents=total_residents,
        total_deceased=total_deceased,
        total_households=total_households,
        total_males=total_males,
        total_females=total_females,
        active_residents=total_residents
    )

@app.route('/user_view_members')
@require_role('user')
def user_view_members():
    """User view - read-only list of all residents"""
    if 'username' not in session:
        return redirect(url_for('login'))
    
    search_query = request.args.get("search", "").strip()
    conn, cursor = get_db()
    sql = """
        SELECT h.*, hh.surname AS household_name
        FROM household h
        LEFT JOIN households hh ON h.household_id = hh.id
    """
    params = []
    if search_query:
        sql += """
            WHERE (
                h.firstname ILIKE %s OR h.middlename ILIKE %s OR h.surname ILIKE %s
                OR CONCAT_WS(' ', h.firstname, h.middlename, h.surname) ILIKE %s
                OR CONCAT_WS(', ', h.surname, h.firstname) ILIKE %s
                OR hh.surname ILIKE %s
            )
        """
        params = [f"%{search_query}%"] * 6
    sql += " ORDER BY h.surname ASC"
    cursor.execute(sql, params)
    members = cursor.fetchall()
    active_residents = sum(1 for member in members if (member.get('status') or 'Active') != 'Deceased')
    deceased_residents = sum(1 for member in members if member.get('status') == 'Deceased')
    conn.close()
    
    return render_template(
        'user_view_members.html',
        members=members,
        keyword=search_query,
        active_residents=active_residents,
        deceased_residents=deceased_residents
    )

@app.route('/user_view_households')
@require_role('user')
def user_view_households():
    """User view - read-only list of all households"""
    if 'username' not in session:
        return redirect(url_for('login'))
    
    search_query = request.args.get("q", "").strip()
    conn, cursor = get_db()
    
    sql = """
        SELECT hh.*, COUNT(h.id) AS member_count 
        FROM households hh
        LEFT JOIN household h ON hh.id = h.household_id
    """
    params = []
    if search_query:
        sql += " WHERE hh.surname ILIKE %s OR hh.house_number ILIKE %s OR hh.address ILIKE %s"
        params = [f"%{search_query}%", f"%{search_query}%", f"%{search_query}%"]
    
    sql += " GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at ORDER BY hh.id DESC"
    cursor.execute(sql, params)
    households = cursor.fetchall()
    conn.close()
    
    return render_template(
        'user_view_households.html',
        households=households,
        keyword=search_query
    )

@app.route('/user_view_household/<int:household_id>')
@require_role('user')
def user_view_household(household_id):
    """User view - read-only household members"""
    if 'username' not in session:
        return redirect(url_for('login'))
    
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM households WHERE id=%s", (household_id,))
    household = cursor.fetchone()
    
    if not household:
        flash("Household not found", "danger")
        return redirect(url_for('user_view_households'))
    
    cursor.execute("SELECT * FROM household WHERE household_id=%s ORDER BY surname ASC", (household_id,))
    members = cursor.fetchall()
    conn.close()
    
    return render_template(
        'user_view_household.html',
        household=household,
        members=members
    )

@app.route('/user_profile')
@require_role('user')
def user_profile():
    """User profile page - can view and edit own profile"""
    if 'username' not in session:
        return redirect(url_for('login'))
    
    username = session.get('username')
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM users WHERE username=%s", (username,))
    user_data = cursor.fetchone()
    conn.close()
    
    return render_template(
        'user_profile.html',
        user_data=user_data,
        user_email=(user_data or {}).get('email'),
        join_date=(user_data or {}).get('signup_date') or 'N/A'
    )

# ===== END USER ROUTES =====

@app.route('/audit-log')
@require_role('admin')
def audit_log():
    if 'username' not in session:
        return redirect(url_for('login'))

    ensure_audit_log_schema()
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM audit_logs ORDER BY created_at DESC")
    logs = cursor.fetchall()

    cursor.execute("""
        SELECT action_type, COUNT(*) AS total
        FROM audit_logs
        GROUP BY action_type
    """)
    action_counts = {row['action_type']: row['total'] for row in cursor.fetchall()}

    cursor.execute("SELECT DISTINCT user_id, username FROM audit_logs ORDER BY username ASC")
    users = cursor.fetchall()
    conn.close()

    return render_template(
        'audit_log.html',
        logs=logs,
        users=users,
        total_logs=len(logs),
        add_count=action_counts.get('ADD', 0),
        update_count=action_counts.get('UPDATE', 0),
        delete_count=action_counts.get('DELETE', 0)
    )

@app.route('/')
@require_role('admin')
def home():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    conn, cursor = get_db()
    debug_log("[DEBUG][statistics] Loading home dashboard totals from household and households tables")
    cursor.execute("""
        SELECT
            COUNT(*) FILTER (WHERE COALESCE(status, 'Active') != 'Deceased') AS total_residents,
            COUNT(*) FILTER (WHERE COALESCE(status, 'Active') = 'Deceased') AS total_deceased,
            COUNT(*) FILTER (WHERE gender = 'Male' AND COALESCE(status, 'Active') != 'Deceased') AS total_males,
            COUNT(*) FILTER (WHERE gender = 'Female' AND COALESCE(status, 'Active') != 'Deceased') AS total_females
        FROM household
    """)
    resident_totals = cursor.fetchone()
    total_residents = resident_totals['total_residents']
    total_deceased = resident_totals['total_deceased']
    total_males = resident_totals['total_males']
    total_females = resident_totals['total_females']
    cursor.execute("SELECT COUNT(*) AS total FROM households")
    total_households = cursor.fetchone()['total']
    debug_log(f"[DEBUG][statistics] totals residents={total_residents} deceased={total_deceased} households={total_households} male={total_males} female={total_females}")
    conn.close()
    
    return render_template(
        'index.html',
        total_residents=total_residents,
        total_deceased=total_deceased,
        total_households=total_households,
        total_males=total_males,
        total_females=total_females
    )

@app.route('/view_members')
@require_role('admin')
def view_members():
    if 'username' not in session:
        return redirect(url_for('login'))
    search_query = request.args.get("search", "").strip()
    conn, cursor = get_db()
    sql = """
        SELECT h.*, hh.surname AS household_name
        FROM household h
        LEFT JOIN households hh ON h.household_id = hh.id
    """
    params = []
    if search_query:
        sql += """
            WHERE (
                h.firstname ILIKE %s OR h.middlename ILIKE %s OR h.surname ILIKE %s
                OR CONCAT_WS(' ', h.firstname, h.middlename, h.surname) ILIKE %s
                OR CONCAT_WS(', ', h.surname, h.firstname) ILIKE %s
                OR hh.surname ILIKE %s
            )
        """
        params = [f"%{search_query}%"] * 6
    sql += " ORDER BY h.surname ASC"
    cursor.execute(sql, params)
    members = cursor.fetchall()
    active_residents = sum(1 for member in members if (member.get('status') or 'Active') != 'Deceased')
    deceased_residents = sum(1 for member in members if member.get('status') == 'Deceased')
    conn.close()
    return render_template('view_members.html', members=members, keyword=search_query, active_residents=active_residents, deceased_residents=deceased_residents)

@app.route('/view_households')
@require_role('admin')
def view_households():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    # Matches the 'q' parameter sent by your script.js
    search_query = request.args.get("q", "").strip()
    conn, cursor = get_db()
    
    # Base SQL to get households and count their members
    sql = """
        SELECT hh.*, COUNT(h.id) AS member_count 
        FROM households hh
        LEFT JOIN household h ON hh.id = h.household_id
    """
    params = []

    # Filter by surname or house number if a search query exists
    if search_query:
        sql += " WHERE hh.surname ILIKE %s OR hh.house_number ILIKE %s OR hh.address ILIKE %s"
        params = [f"%{search_query}%", f"%{search_query}%", f"%{search_query}%"]

    sql += " GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at ORDER BY hh.id DESC"
    
    cursor.execute(sql, params)
    households = cursor.fetchall()
    conn.close()
    
    return render_template('view_households.html', households=households, keyword=search_query)

@app.route('/print_reports')
@require_role('admin')
def print_reports():
    if 'username' not in session:
        return redirect(url_for('login'))

    ensure_audit_log_schema()
    conn, cursor = get_db()
    cursor.execute("SELECT id, surname, house_number, address FROM households ORDER BY surname ASC")
    households = cursor.fetchall()
    cursor.execute("SELECT DISTINCT username FROM audit_logs ORDER BY username ASC")
    audit_users = cursor.fetchall()
    conn.close()

    return render_template('print_reports.html', households=households, audit_users=audit_users)

@app.route('/search_households')
@require_role('admin', 'user')
def search_households():
    q = request.args.get("q", "").strip()
    if not q: return jsonify([])
    
    conn = None
    try:
        conn, cursor = get_db()
        debug_log(f"[DEBUG][search] /search_households q={q!r}")
        # Updated query to join and get member counts for the suggestion dropdown
        query = """
            SELECT hh.id, hh.surname, hh.house_number, COUNT(h.id) as count
            FROM households hh
            LEFT JOIN household h ON hh.id = h.household_id
            WHERE hh.surname ILIKE %s
               OR CAST(hh.id AS TEXT) ILIKE %s
               OR hh.house_number ILIKE %s
               OR hh.address ILIKE %s
               OR CONCAT_WS(' ', hh.surname, hh.house_number, hh.address) ILIKE %s
            GROUP BY hh.id, hh.surname, hh.house_number, hh.address
            ORDER BY hh.surname ASC
            LIMIT 10
        """
        cursor.execute(query, (f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"))
        res = cursor.fetchall()

        return jsonify([{
            "id": r['id'], 
            "name": r['surname'], 
            "house_number": r['house_number'],
            "subtext": f"House No: {r['house_number'] or 'N/A'} | Family Members: {r['count']}",
            "type": "household"
        } for r in res])
    except Exception as e:
        print(f"[ERROR][search] /search_households error: {e}")
        return jsonify([]), 500
    finally:
        if conn:
            conn.close()

@app.route('/search_members')
@require_role('admin', 'user')
def search_members():
    q = request.args.get("q", "").strip()
    if not q: return jsonify([])
    
    conn = None
    try:
        conn, cursor = get_db()
        debug_log(f"[DEBUG][search] /search_members q={q!r}")
        # This query pulls the surname from the 'households' table
        query = """
            SELECT h.id, h.firstname, h.middlename, h.surname, h.occupation, hh.surname as hh_name
            FROM household h
            LEFT JOIN households hh ON h.household_id = hh.id
            WHERE h.firstname ILIKE %s
               OR h.middlename ILIKE %s
               OR h.surname ILIKE %s
               OR CONCAT_WS(' ', h.firstname, h.middlename, h.surname) ILIKE %s
               OR CONCAT_WS(', ', h.surname, h.firstname) ILIKE %s
               OR hh.surname ILIKE %s
            ORDER BY h.surname ASC, h.firstname ASC
            LIMIT 10
        """
        cursor.execute(query, (f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"))
        res = cursor.fetchall()

        formatted_results = []
        for row in res:
            formatted_results.append({
                "id": row['id'],
                "name": f"{row['surname']}, {row['firstname']}",
                "household_name": row['hh_name'],
                "occupation": row.get('occupation') or "Resident",
                "type": "resident"
            })
        return jsonify(formatted_results)
    except Exception as e:
        print(f"[ERROR][search] /search_members error: {e}")
        return jsonify([]), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/search_residents')
@require_role('admin', 'user')
def api_search_residents():
    q = request.args.get("q", "").strip()
    if not q:
        return jsonify([])

    conn = None
    try:
        conn, cursor = get_db()
        debug_log(f"[DEBUG][search] /api/search_residents q={q!r}")
        cursor.execute("""
            SELECT h.id, h.firstname, h.middlename, h.surname, h.gender, h.status,
                   hh.id AS household_id, hh.surname AS household_name, hh.house_number
            FROM household h
            LEFT JOIN households hh ON h.household_id = hh.id
            WHERE h.firstname ILIKE %s
               OR h.middlename ILIKE %s
               OR h.surname ILIKE %s
               OR CONCAT_WS(' ', h.firstname, h.middlename, h.surname) ILIKE %s
               OR CONCAT_WS(', ', h.surname, h.firstname) ILIKE %s
               OR hh.surname ILIKE %s
            ORDER BY h.surname ASC
            LIMIT 10
        """, (f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"))
        residents = cursor.fetchall()

        cursor.execute("""
            SELECT hh.id, hh.surname, hh.house_number, COUNT(h.id) AS member_count
            FROM households hh
            LEFT JOIN household h ON hh.id = h.household_id
            WHERE hh.surname ILIKE %s OR hh.house_number ILIKE %s OR hh.address ILIKE %s
            GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at
            ORDER BY hh.surname ASC
            LIMIT 10
        """, (f"%{q}%", f"%{q}%", f"%{q}%"))
        households = cursor.fetchall()

        results = []
        user_role = session.get('role')
        for resident in residents:
            full_name = " ".join(filter(None, [
                resident.get('firstname'),
                resident.get('middlename'),
                resident.get('surname')
            ]))
            if resident.get('household_id'):
                resident_url = url_for(
                    'user_view_household' if user_role == 'user' else 'view_household',
                    household_id=resident.get('household_id')
                )
            else:
                resident_url = url_for('user_view_members' if user_role == 'user' else 'view_members')
            results.append({
                'type': 'Resident',
                'name': full_name,
                'details': f"Household: {resident.get('household_name') or 'Unassigned'}",
                'status': resident.get('status') or 'Active',
                'id': resident.get('id'),
                'url': resident_url
            })

        for household in households:
            results.append({
                'type': 'Household',
                'name': household.get('surname'),
                'details': f"House No: {household.get('house_number') or 'N/A'} | Members: {household.get('member_count') or 0}",
                'status': 'Read-only',
                'id': household.get('id'),
                'url': url_for(
                    'user_view_household' if user_role == 'user' else 'view_household',
                    household_id=household.get('id')
                )
            })

        return jsonify(results)
    except Exception as e:
        print(f"[ERROR][search] /api/search_residents error: {e}")
        return jsonify([]), 500
    finally:
        if conn:
            conn.close()



@app.route('/add_member', methods=['GET', 'POST'])
@require_role('admin')
def add_member():
    if 'username' not in session:
        return redirect(url_for('login'))
    household_id_from_url = request.args.get('household_id')
    conn, cursor = get_db()
    if request.method == 'POST':
        data = (request.form['firstname'].strip(), request.form['middlename'].strip(),
                request.form['surname'].strip(), request.form['age'],
                request.form['birthdate'], request.form['gender'],
                request.form['civil_status'], request.form['occupation'],
                request.form['household_id'])
        fixed_hh_id = request.form.get('household_id_fixed')
        try:
            cursor.execute("""
                INSERT INTO household (firstname, middlename, surname, age, birthdate, gender, civil_status, occupation, household_id, status) 
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'Active')
                RETURNING id
            """, data)
            new_member_id = cursor.fetchone()['id']
            conn.commit()
            full_name = f"{request.form['firstname']} {request.form['surname']}"
            log_audit(
                session.get('username', 'admin'),
                'ADD',
                f'Added resident profile for {full_name}',
                request.form.get('household_id', 'N/A'),
                user_id=session.get('user_id'),
                target_type='Resident',
                target_id=new_member_id,
                new_value={
                    'firstname': request.form['firstname'].strip(),
                    'middlename': request.form['middlename'].strip(),
                    'surname': request.form['surname'].strip(),
                    'age': request.form['age'],
                    'birthdate': request.form['birthdate'],
                    'gender': request.form['gender'],
                    'civil_status': request.form['civil_status'],
                    'occupation': request.form['occupation'],
                    'household_id': request.form['household_id'],
                    'status': 'Active'
                }
            )
            conn.close()
            return redirect(url_for('view_household', household_id=fixed_hh_id) if fixed_hh_id else url_for('view_members'))
        except Exception as e:
            conn.rollback()
            cursor.execute("SELECT * FROM households")
            households_list = cursor.fetchall()
            conn.close()
            return render_template('add_member.html', households=households_list, error="Database error or duplicate record.")
    cursor.execute("SELECT * FROM households")
    households_list = cursor.fetchall()
    prefill = {k: request.args.get(v, '') for k, v in {'firstname':'fname', 'middlename':'mname', 'surname':'sname', 'age':'age', 'birthdate':'birthdate', 'gender':'gender', 'civil_status':'civil_status', 'occupation':'occupation'}.items()}
    household_name = ""
    if household_id_from_url:
        cursor.execute("SELECT surname FROM households WHERE id=%s", (household_id_from_url,))
        res = cursor.fetchone()
        if res:
            prefill['surname'] = res['surname']
            household_name = res['surname']
    conn.close()
    return render_template('add_member.html', households=households_list, prefill=prefill, household_id=household_id_from_url, household_name=household_name)

@app.route('/edit_member/<int:id>', methods=['GET', 'POST'])
@require_role('admin')
def edit_member(id):
    if 'username' not in session:
        return redirect(url_for('login'))
    conn, cursor = get_db()
    if request.method == 'POST':
        cursor.execute("SELECT * FROM household WHERE id=%s", (id,))
        old_member = cursor.fetchone()
        data = (request.form.get('firstname'), request.form.get('middlename', ''), request.form.get('surname'),
                request.form.get('age'), request.form.get('birthdate'), request.form.get('gender'),
                request.form.get('civil_status'), request.form.get('occupation', ''),
                request.form.get('household_id'), request.form.get('status', 'Active'), id)
        fixed_hh_id = request.form.get('household_id_fixed') or request.args.get('household_id')
        try:
            cursor.execute("""
                UPDATE household SET firstname=%s, middlename=%s, surname=%s, age=%s, birthdate=%s,
                gender=%s, civil_status=%s, occupation=%s, household_id=%s, status=%s WHERE id=%s
            """, data)
            conn.commit()
            full_name = f"{request.form.get('firstname')} {request.form.get('surname')}"
            log_audit(
                session.get('username', 'admin'),
                'UPDATE',
                f'Updated resident profile for {full_name}',
                request.form.get('household_id', 'N/A'),
                user_id=session.get('user_id'),
                target_type='Resident',
                target_id=id,
                old_value=old_member,
                new_value={
                    'firstname': request.form.get('firstname'),
                    'middlename': request.form.get('middlename', ''),
                    'surname': request.form.get('surname'),
                    'age': request.form.get('age'),
                    'birthdate': request.form.get('birthdate'),
                    'gender': request.form.get('gender'),
                    'civil_status': request.form.get('civil_status'),
                    'occupation': request.form.get('occupation', ''),
                    'household_id': request.form.get('household_id'),
                    'status': request.form.get('status', 'Active')
                }
            )
            conn.close()
            return redirect(url_for('view_household', household_id=fixed_hh_id) if fixed_hh_id else url_for('view_members'))
        except Exception as e:
            conn.rollback()
            flash(f"Error: {str(e)}", "danger")
    cursor.execute("SELECT * FROM household WHERE id=%s", (id,))
    member = cursor.fetchone()
    cursor.execute("SELECT * FROM households")
    households = cursor.fetchall()
    conn.close()
    return render_template('edit_member.html', member=member, households=households)

@app.route('/delete_member/<int:id>')
@require_role('admin')
def delete_member(id):
    if 'username' not in session:
        return redirect(url_for('login'))
    household_id = request.args.get('household_id')
    conn, cursor = get_db()
    # Get member info before deleting for audit log
    cursor.execute("SELECT * FROM household WHERE id=%s", (id,))
    member = cursor.fetchone()
    cursor.execute("DELETE FROM household WHERE id=%s", (id,))
    conn.commit()
    # Log the action
    if member:
        full_name = f"{member.get('firstname')} {member.get('surname')}"
        log_audit(
            session.get('username', 'admin'),
            'DELETE',
            f'Deleted resident profile for {full_name}',
            household_id or member.get('household_id') or 'N/A',
            user_id=session.get('user_id'),
            target_type='Resident',
            target_id=id,
            old_value=member
        )
    conn.close()
    return redirect(url_for('view_household', household_id=household_id) if household_id else url_for('view_members'))


@app.route('/add_household', methods=['GET', 'POST'])
@require_role('admin')
def add_household():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    from_member = request.args.get('from_member', 'false')
    
    # 1. CAPTURE ALL RESIDENT DATA FROM URL
    # This ensures if you clicked "Register Household" from the member form, 
    # your typed data is stored in memory here.
    resident_data = {
        'fname': request.args.get('fname', ''),
        'mname': request.args.get('mname', ''),
        'sname': request.args.get('sname', ''),
        'age': request.args.get('age', ''),
        'birthdate': request.args.get('birthdate', ''),
        'gender': request.args.get('gender', ''),
        'civil_status': request.args.get('civil_status', ''),
        'occupation': request.args.get('occupation', '')
    }

    if request.method == 'POST':
        surname = request.form['surname'].strip()
        h_num = request.form.get('house_number', '').strip()
        addr = request.form.get('address', '').strip()
        
        # Capture the hidden resident data from the POST form as well
        # (This is needed if the household already exists and we need to re-render the page)
        post_resident_data = {
            'fname': request.form.get('fname', ''),
            'mname': request.form.get('mname', ''),
            'sname': request.form.get('sname', ''),
            'age': request.form.get('age', ''),
            'birthdate': request.form.get('birthdate', ''),
            'gender': request.form.get('gender', ''),
            'civil_status': request.form.get('civil_status', ''),
            'occupation': request.form.get('occupation', '')
        }

        conn, cursor = get_db()
        cursor.execute("SELECT * FROM households WHERE surname=%s AND house_number=%s", (surname, h_num))
        
        if cursor.fetchone():
            conn.close()
            # 2. PASS DATA BACK ON ERROR
            # If the household exists, we must pass the resident data back so the hidden inputs stay filled
            return render_template('add_household.html', 
                                 error="Household already exists.", 
                                 from_member=from_member, 
                                 **post_resident_data)
        
        cursor.execute("INSERT INTO households (surname, house_number, address) VALUES (%s, %s, %s) RETURNING id", (surname, h_num, addr))
        new_household_id = cursor.fetchone()['id']
        conn.commit()
        log_audit(
            session.get('username', 'admin'),
            'ADD',
            f'Added household record for {surname}',
            new_household_id,
            user_id=session.get('user_id'),
            target_type='Household',
            target_id=new_household_id,
            new_value={
                'surname': surname,
                'house_number': h_num,
                'address': addr
            }
        )
        conn.close()

        # 3. REDIRECT WITH ALL DATA
        # Instead of just sname, we pass the whole dictionary back to add_member
        if from_member == 'true':
            return redirect(url_for('add_member', **post_resident_data))
        
        return redirect(url_for('view_households'))

    # 4. PASS DATA ON INITIAL GET
    return render_template('add_household.html', from_member=from_member, **resident_data)
@app.route('/view_household/<int:household_id>')
@require_role('admin')
def view_household(household_id):
    if 'username' not in session:
        return redirect(url_for('login'))
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM households WHERE id=%s", (household_id,))
    household = cursor.fetchone()
    cursor.execute("SELECT * FROM household WHERE household_id = %s", (household_id,))
    members = cursor.fetchall()
    conn.close()
    return render_template('view_household.html', household=household, members=members)

@app.route('/delete_household/<int:id>')
@require_role('admin')
def delete_household(id):
    if 'username' not in session:
        return redirect(url_for('login'))
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM households WHERE id=%s", (id,))
    household = cursor.fetchone()
    cursor.execute("SELECT * FROM household WHERE household_id=%s", (id,))
    members = cursor.fetchall()
    cursor.execute("DELETE FROM household WHERE household_id=%s", (id,))
    cursor.execute("DELETE FROM households WHERE id=%s", (id,))
    conn.commit()
    if household:
        log_audit(
            session.get('username', 'admin'),
            'DELETE',
            f"Deleted household record for {household.get('surname')}",
            id,
            user_id=session.get('user_id'),
            target_type='Household',
            target_id=id,
            old_value={
                'household': household,
                'members': members
            }
        )
    conn.close()
    return redirect(url_for('view_households'))

@app.route('/print_household_members/<int:household_id>')
@require_role('admin')
def print_household_members(household_id):
    if 'username' not in session:
        return redirect(url_for('login'))
        
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM households WHERE id=%s", (household_id,))
    household = cursor.fetchone()
    cursor.execute("SELECT * FROM household WHERE household_id=%s ORDER BY birthdate ASC", (household_id,))
    members = cursor.fetchall()
    conn.close()

    doc = Document()
    
    # --- HEADER ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = header.add_run("BARANGAY HOUSEHOLD INFORMATION RECORD")
    title.bold = True
    title.font.size = Pt(14)
    
    # --- FORMAL DATA BOX ---
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info = doc.add_paragraph()
    info.add_run(f"HOUSEHOLD SURNAME: ").bold = True
    info.add_run(f"{household['surname'].upper()}\n")
    info.add_run(f"HOUSE NUMBER: ").bold = True
    info.add_run(f"{household['house_number']}\n")
    info.add_run(f"COMPLETE ADDRESS: ").bold = True
    info.add_run(f"{household['address']}")
    
    doc.add_paragraph("\nLIST OF REGISTERED MEMBERS:").runs[0].bold = True

    # --- TABLE ---
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(['Full Name', 'Birthdate', 'Civil Status', 'Occupation', 'Status']):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for m in members:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{m['surname']}, {m['firstname']}"
        row_cells[1].text = str(m['birthdate'])
        row_cells[2].text = str(m['civil_status'])
        row_cells[3].text = str(m['occupation'] or 'N/A')
        row_cells[4].text = str(m['status'] or 'Active')

    doc.add_paragraph(f"\nTotal Registered Members: {len(members)}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Certified Correct: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name=f"Household_{household['surname']}.docx")

@app.route('/print_households_report')
@require_role('admin')
def print_households_report():
    if 'username' not in session:
        return redirect(url_for('login'))

    household_id = request.args.get('household_id', '').strip()
    search = request.args.get('search', '').strip()
    gender = request.args.get('gender', '').strip()
    status = request.args.get('status', '').strip()
    month = request.args.get('month', '').strip()

    if gender not in {'Male', 'Female'}:
        gender = ''
    if status not in {'Active', 'Deceased'}:
        status = ''
    try:
        month_int = int(month) if month else None
        if month == 'year':
            month_int = 'year'
        elif month_int and (month_int < 1 or month_int > 12):
            month_int = None
    except ValueError:
        month_int = None

    ensure_audit_log_schema()
    conn, cursor = get_db()
    query = """
        SELECT *
        FROM (
            SELECT
                hh.id,
                hh.surname,
                hh.house_number,
                hh.address,
                COUNT(h.id) AS total_members,
                SUM(CASE WHEN h.gender = 'Male' THEN 1 ELSE 0 END) AS male_members,
                SUM(CASE WHEN h.gender = 'Female' THEN 1 ELSE 0 END) AS female_members,
                SUM(CASE WHEN COALESCE(h.status, 'Active') != 'Deceased' THEN 1 ELSE 0 END) AS active_members,
                SUM(CASE WHEN COALESCE(h.status, 'Active') = 'Deceased' THEN 1 ELSE 0 END) AS deceased_members,
                COALESCE((
                    SELECT MIN(al.created_at)
                    FROM audit_logs al
                    WHERE al.target_type = 'Household'
                      AND al.target_id = CAST(hh.id AS TEXT)
                      AND al.action_type = 'ADD'
                ), hh.created_at) AS household_registered
            FROM households hh
            LEFT JOIN household h ON hh.id = h.household_id
            GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at
        ) household_report
        WHERE 1=1
    """
    params = []

    if household_id:
        query += " AND id = %s"
        params.append(household_id)
    if search:
        query += " AND (surname ILIKE %s OR house_number ILIKE %s OR address ILIKE %s)"
        params.extend([f"%{search}%"] * 3)
    if gender == 'Male':
        query += " AND male_members > 0"
    elif gender == 'Female':
        query += " AND female_members > 0"
    if status == 'Active':
        query += " AND active_members > 0"
    elif status == 'Deceased':
        query += " AND deceased_members > 0"
    if month_int == 'year':
        query += f" AND {_current_year_filter_sql('household_registered')}"
    elif month_int:
        query += f" AND {_month_filter_sql('household_registered')}"
        params.append(month_int)

    query += " ORDER BY surname ASC"
    print(f"[DEBUG][print] print_households_report filters household_id={household_id!r} search={search!r} gender={gender!r} status={status!r} month={month!r}")
    cursor.execute(query, params)
    households = cursor.fetchall()
    conn.close()

    month_names = {
        1: 'January', 2: 'February', 3: 'March', 4: 'April',
        5: 'May', 6: 'June', 7: 'July', 8: 'August',
        9: 'September', 10: 'October', 11: 'November', 12: 'December'
    }

    def fmt_doc_date(value):
        if not value:
            return 'N/A'
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace('Z', '+00:00'))
            except ValueError:
                return value
        return value.strftime('%b %d, %Y') if hasattr(value, 'strftime') else str(value)

    filters = []
    if household_id:
        filters.append(f"Household ID: {household_id}")
    if search:
        filters.append(f"Search: {search}")
    if gender:
        filters.append(f"Resident Gender Included: {gender}")
    if status:
        filters.append(f"Resident Status Included: {status}")
    if month_int == 'year':
        filters.append(f"Household Registration: This Year")
    elif month_int:
        filters.append(f"Household Registration Month: {month_names[month_int]}")
    if not filters:
        filters.append("Filters: All households")

    doc = Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("REPUBLIC OF THE PHILIPPINES\nPROVINCE OF ORIENTAL MINDORO\nMUNICIPALITY OF BONGABONG\nBARANGAY REGISTRY OFFICE")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_heading('HOUSEHOLDS REPORT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in filters:
        doc.add_paragraph(line)
    doc.add_paragraph(f"Total Records: {len(households)}")

    headers = ['No.', 'House No.', 'Household', 'Address', 'Total', 'Male', 'Female', 'Active', 'Deceased', 'Registered']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for idx, h in enumerate(households, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(h.get('house_number') or 'N/A')
        row[2].text = str(h.get('surname') or 'N/A')
        row[3].text = str(h.get('address') or 'N/A')
        row[4].text = str(h.get('total_members') or 0)
        row[5].text = str(h.get('male_members') or 0)
        row[6].text = str(h.get('female_members') or 0)
        row[7].text = str(h.get('active_members') or 0)
        row[8].text = str(h.get('deceased_members') or 0)
        row[9].text = fmt_doc_date(h.get('household_registered'))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name="Barangay_Households_Report.docx")

@app.route('/print_audit_logs_report')
@require_role('admin')
def print_audit_logs_report():
    if 'username' not in session:
        return redirect(url_for('login'))

    action_type = request.args.get('action_type', '').strip()
    target_type = request.args.get('target_type', '').strip()
    username = request.args.get('username', '').strip()
    month = request.args.get('month', '').strip()

    if action_type not in {'ADD', 'UPDATE', 'DELETE'}:
        action_type = ''
    if target_type not in {'Resident', 'Household', 'System'}:
        target_type = ''
    try:
        month_int = int(month) if month else None
        if month == 'year':
            month_int = 'year'
        elif month_int and (month_int < 1 or month_int > 12):
            month_int = None
    except ValueError:
        month_int = None

    ensure_audit_log_schema()
    conn, cursor = get_db()
    query = """
        SELECT id, username, action_type, target_type, target_id, details,
               household_context, status, created_at
        FROM audit_logs
        WHERE 1=1
    """
    params = []

    if action_type:
        query += " AND action_type = %s"
        params.append(action_type)
    if target_type:
        query += " AND target_type = %s"
        params.append(target_type)
    if username:
        query += " AND username = %s"
        params.append(username)
    if month_int == 'year':
        query += f" AND {_current_year_filter_sql('created_at')}"
    elif month_int:
        query += f" AND {_month_filter_sql('created_at')}"
        params.append(month_int)

    query += " ORDER BY created_at DESC"
    print(f"[DEBUG][print] print_audit_logs_report filters action_type={action_type!r} target_type={target_type!r} username={username!r} month={month!r}")
    cursor.execute(query, params)
    logs = cursor.fetchall()
    conn.close()

    month_names = {
        1: 'January', 2: 'February', 3: 'March', 4: 'April',
        5: 'May', 6: 'June', 7: 'July', 8: 'August',
        9: 'September', 10: 'October', 11: 'November', 12: 'December'
    }

    def fmt_doc_date(value):
        if not value:
            return 'N/A'
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace('Z', '+00:00'))
            except ValueError:
                return value
        return value.strftime('%b %d, %Y %I:%M %p') if hasattr(value, 'strftime') else str(value)

    filters = []
    if action_type:
        filters.append(f"Action: {action_type}")
    if target_type:
        filters.append(f"Target: {target_type}")
    if username:
        filters.append(f"User: {username}")
    if month_int == 'year':
        filters.append(f"Period: This Year")
    elif month_int:
        filters.append(f"Month: {month_names[month_int]}")
    if not filters:
        filters.append("Filters: All audit log entries")

    doc = Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("REPUBLIC OF THE PHILIPPINES\nPROVINCE OF ORIENTAL MINDORO\nMUNICIPALITY OF BONGABONG\nBARANGAY REGISTRY OFFICE")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_heading('AUDIT LOG REPORT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in filters:
        doc.add_paragraph(line)
    doc.add_paragraph(f"Total Records: {len(logs)}")

    headers = ['No.', 'Date / Time', 'User', 'Action', 'Target', 'Target ID', 'Details', 'Status']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for idx, log in enumerate(logs, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = fmt_doc_date(log.get('created_at'))
        row[2].text = str(log.get('username') or 'N/A')
        row[3].text = str(log.get('action_type') or 'N/A')
        row[4].text = str(log.get('target_type') or 'N/A')
        row[5].text = str(log.get('target_id') or 'N/A')
        row[6].text = str(log.get('details') or 'N/A')
        row[7].text = str(log.get('status') or 'SUCCESS')

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name="Barangay_Audit_Log_Report.docx")

@app.route('/print_all_members')
@require_role('admin')
def print_all_members():
    if 'username' not in session:
        return redirect(url_for('login'))

    report_type = request.args.get('report_type', 'all').strip()
    gender = request.args.get('gender', '').strip()
    household_id = request.args.get('household_id', '').strip()
    month = request.args.get('month', '').strip()

    if report_type not in {'all', 'registered', 'deceased'}:
        report_type = 'all'
    if gender not in {'Male', 'Female'}:
        gender = ''
    try:
        month_int = int(month) if month else None
        if month == 'year':
            month_int = 'year'
        elif month_int and (month_int < 1 or month_int > 12):
            month_int = None
    except ValueError:
        month_int = None

    ensure_audit_log_schema()
    conn, cursor = get_db()
    query = """
        SELECT *
        FROM (
            SELECT
                h.*,
                hh.surname AS household_name,
                COALESCE((
                    SELECT MIN(al.created_at)
                    FROM audit_logs al
                    WHERE al.target_type = 'Resident'
                      AND al.target_id = CAST(h.id AS TEXT)
                      AND al.action_type = 'ADD'
                ), h.created_at) AS registration_date,
                (
                    SELECT MIN(al.created_at)
                    FROM audit_logs al
                    WHERE al.target_type = 'Resident'
                      AND al.target_id = CAST(h.id AS TEXT)
                      AND al.action_type = 'UPDATE'
                      AND (
                          al.new_value ILIKE '%"status": "Deceased"%'
                          OR al.new_value ILIKE '%"status":"Deceased"%'
                      )
                ) AS deceased_date
            FROM household h
            LEFT JOIN households hh ON h.household_id = hh.id
        ) resident_report
        WHERE 1=1
    """
    params = []

    if report_type == 'deceased':
        query += " AND COALESCE(status, 'Active') = 'Deceased'"
    elif report_type == 'registered':
        query += " AND registration_date IS NOT NULL"

    if gender:
        query += " AND gender = %s"
        params.append(gender)
    if household_id:
        query += " AND household_id = %s"
        params.append(household_id)

    date_column = 'deceased_date' if report_type == 'deceased' else 'registration_date'
    if report_type in {'registered', 'deceased'}:
        if month_int == 'year':
            query += f" AND {_current_year_filter_sql(date_column)}"
        elif month_int:
            query += f" AND {_month_filter_sql(date_column)}"
            params.append(month_int)

    query += " ORDER BY surname ASC, firstname ASC"
    print(f"[DEBUG][print] print_all_members filters report_type={report_type!r} gender={gender!r} household_id={household_id!r} month={month!r}")
    cursor.execute(query, params)
    members = cursor.fetchall()
    conn.close()

    month_names = {
        1: 'January', 2: 'February', 3: 'March', 4: 'April',
        5: 'May', 6: 'June', 7: 'July', 8: 'August',
        9: 'September', 10: 'October', 11: 'November', 12: 'December'
    }

    def fmt_doc_date(value):
        if not value:
            return 'N/A'
        if isinstance(value, str):
            try:
                value = datetime.fromisoformat(value.replace('Z', '+00:00'))
            except ValueError:
                return value
        return value.strftime('%b %d, %Y') if hasattr(value, 'strftime') else str(value)

    title_map = {
        'all': 'RESIDENTS MASTER LIST',
        'registered': 'REGISTERED RESIDENTS REPORT',
        'deceased': 'DECEASED RESIDENTS REPORT'
    }
    filter_lines = []
    if gender:
        filter_lines.append(f"Gender: {gender}")
    if household_id:
        filter_lines.append(f"Household ID: {household_id}")
    if report_type in {'registered', 'deceased'}:
        date_label = 'Death / Deceased Update Date' if report_type == 'deceased' else 'Registration Date'
        if month_int == 'year':
            filter_lines.append(f"{date_label}: This Year")
        elif month_int:
            filter_lines.append(f"{date_label} Month: {month_names[month_int]}")
    if not filter_lines:
        filter_lines.append('Filters: All residents')

    doc = Document()
    
    # --- SETUP LANDSCAPE FOR MORE SPACE ---
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # --- FORMAL HEADER ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("REPUBLIC OF THE PHILIPPINES\nPROVINCE OF ORIENTAL MINDORO\nMUNICIPALITY OF BONGABONG\nBARANGAY REGISTRY OFFICE")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_heading(title_map[report_type], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for line in filter_lines:
        doc.add_paragraph(line)
    doc.add_paragraph(f"Total Records: {len(members)}")

    # --- DATA TABLE ---
    headers = ['No.', 'Full Name', 'Birthdate', 'Age', 'Gender', 'Occupation', 'Status']
    if report_type in {'all', 'registered'}:
        headers.append('Registered Date')
    if report_type in {'all', 'deceased'}:
        headers.append('Death / Deceased Update Date')

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for idx, m in enumerate(members, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"{m['surname']}, {m['firstname']} {m['middlename'] or ''}"
        row_cells[2].text = str(m['birthdate']) if m['birthdate'] else 'N/A'
        row_cells[3].text = str(m['age'])
        row_cells[4].text = str(m['gender'])
        row_cells[5].text = str(m['occupation'] or 'N/A')
        row_cells[6].text = str(m['status'] or 'Active')
        col = 7
        if report_type in {'all', 'registered'}:
            row_cells[col].text = fmt_doc_date(m.get('registration_date'))
            col += 1
        if report_type in {'all', 'deceased'}:
            row_cells[col].text = fmt_doc_date(m.get('deceased_date')) if (m.get('status') == 'Deceased') else 'N/A'

    # --- SPACED SIGNATURE SECTION ---
    doc.add_paragraph("\n" * 3) # Extra space before signatures
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.width = Inches(9) # Make signature table wide
    
    sig_table.rows[0].cells[0].text = "Prepared by:"
    sig_table.rows[0].cells[1].text = "Attested by:"
    
    # Add actual lines and titles
    p1 = sig_table.rows[1].cells[0].paragraphs[0]
    p1.add_run("\n\n__________________________\nBarangay Secretary").bold = True
    
    p2 = sig_table.rows[1].cells[1].paragraphs[0]
    p2.add_run("\n\n__________________________\nBarangay Captain").bold = True

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename_parts = ['Barangay', title_map[report_type].title().replace(' ', '_')]
    if gender:
        filename_parts.append(gender)
    if month_int:
        filename_parts.append(month_names[month_int])
    return send_file(file_stream, as_attachment=True, download_name=f"{'_'.join(filename_parts)}.docx")

@app.route('/api/preview/residents')
@require_role('admin')
def api_preview_residents():
    """Preview API for resident reports"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    conn = None
    try:
        report_type = request.args.get('report_type', 'all').strip()
        gender = request.args.get('gender', '').strip()
        household_id = request.args.get('household_id', '').strip()
        month = request.args.get('month', '').strip()

        if report_type not in {'all', 'registered', 'deceased'}:
            report_type = 'all'
        if gender not in {'Male', 'Female'}:
            gender = ''
        try:
            month_int = int(month) if month else None
            if month == 'year':
                month_int = 'year'
            elif month_int and (month_int < 1 or month_int > 12):
                month_int = None
        except ValueError:
            month_int = None

        ensure_audit_log_schema()
        conn, cursor = get_db()
        query = """
            SELECT *
            FROM (
                SELECT
                    h.*,
                    hh.surname AS household_name,
                    COALESCE((
                        SELECT MIN(al.created_at)
                        FROM audit_logs al
                        WHERE al.target_type = 'Resident'
                          AND al.target_id = CAST(h.id AS TEXT)
                          AND al.action_type = 'ADD'
                    ), h.created_at) AS registration_date,
                    (
                        SELECT MIN(al.created_at)
                        FROM audit_logs al
                        WHERE al.target_type = 'Resident'
                          AND al.target_id = CAST(h.id AS TEXT)
                          AND al.action_type = 'UPDATE'
                          AND (
                              al.new_value ILIKE '%"status": "Deceased"%'
                              OR al.new_value ILIKE '%"status":"Deceased"%'
                          )
                    ) AS deceased_date
                FROM household h
                LEFT JOIN households hh ON h.household_id = hh.id
            ) resident_report
            WHERE 1=1
        """
        params = []

        if report_type == 'deceased':
            query += " AND COALESCE(status, 'Active') = 'Deceased'"
        elif report_type == 'registered':
            query += " AND registration_date IS NOT NULL"

        if gender:
            query += " AND gender = %s"
            params.append(gender)
        if household_id:
            query += " AND household_id = %s"
            params.append(household_id)

        date_column = 'deceased_date' if report_type == 'deceased' else 'registration_date'
        if report_type in {'registered', 'deceased'}:
            if month_int == 'year':
                query += f" AND {_current_year_filter_sql(date_column)}"
            elif month_int:
                query += f" AND {_month_filter_sql(date_column)}"
                params.append(month_int)

        query += " ORDER BY surname ASC, firstname ASC LIMIT 100"
        debug_log(f"[DEBUG][preview] /api/preview/residents filters report_type={report_type!r} gender={gender!r} household_id={household_id!r} month={month!r}")
        cursor.execute(query, params)
        members = cursor.fetchall()

        # Format for preview
        items = []
        for m in members:
            items.append({
                'id': m.get('id'),
                'full_name': f"{m.get('surname', '')} {m.get('firstname', '')}".strip(),
                'age': m.get('age'),
                'gender': m.get('gender'),
                'status': m.get('status', 'Active'),
                'household_name': m.get('household_name', 'N/A'),
                'registration_date': str(m.get('registration_date')) if m.get('registration_date') else None
            })

        return jsonify({
            'total': len(items),
            'items': items
        })
    except Exception as e:
        print(f"[ERROR][preview] /api/preview/residents error: {e}")
        return jsonify({'error': str(e), 'total': 0, 'items': []}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/preview/households')
@require_role('admin')
def api_preview_households():
    """Preview API for household reports"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    conn = None
    try:
        household_id = request.args.get('household_id', '').strip()
        search = request.args.get('search', '').strip()
        gender = request.args.get('gender', '').strip()
        status = request.args.get('status', '').strip()
        month = request.args.get('month', '').strip()

        if gender not in {'Male', 'Female'}:
            gender = ''
        if status not in {'Active', 'Deceased'}:
            status = ''
        try:
            month_int = int(month) if month else None
            if month == 'year':
                month_int = 'year'
            elif month_int and (month_int < 1 or month_int > 12):
                month_int = None
        except ValueError:
            month_int = None

        ensure_audit_log_schema()
        conn, cursor = get_db()
        query = """
            SELECT *
            FROM (
                SELECT
                    hh.id,
                    hh.surname,
                    hh.house_number,
                    hh.address,
                    COUNT(h.id) AS member_count,
                    SUM(CASE WHEN h.gender = 'Male' THEN 1 ELSE 0 END) AS male_members,
                    SUM(CASE WHEN h.gender = 'Female' THEN 1 ELSE 0 END) AS female_members,
                    SUM(CASE WHEN COALESCE(h.status, 'Active') != 'Deceased' THEN 1 ELSE 0 END) AS active_count,
                    SUM(CASE WHEN COALESCE(h.status, 'Active') = 'Deceased' THEN 1 ELSE 0 END) AS deceased_count,
                    COALESCE((
                        SELECT MIN(al.created_at)
                        FROM audit_logs al
                        WHERE al.target_type = 'Household'
                          AND al.target_id = CAST(hh.id AS TEXT)
                          AND al.action_type = 'ADD'
                    ), hh.created_at) AS registration_date
                FROM households hh
                LEFT JOIN household h ON hh.id = h.household_id
                GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at
            ) household_report
            WHERE 1=1
        """
        params = []

        if household_id:
            query += " AND id = %s"
            params.append(household_id)
        if search:
            query += " AND (surname ILIKE %s OR house_number ILIKE %s OR address ILIKE %s)"
            params.extend([f"%{search}%"] * 3)
        if gender == 'Male':
            query += " AND male_members > 0"
        elif gender == 'Female':
            query += " AND female_members > 0"
        if status == 'Active':
            query += " AND active_count > 0"
        elif status == 'Deceased':
            query += " AND deceased_count > 0"
        if month_int == 'year':
            query += f" AND {_current_year_filter_sql('registration_date')}"
        elif month_int:
            query += f" AND {_month_filter_sql('registration_date')}"
            params.append(month_int)

        query += " ORDER BY surname ASC LIMIT 100"
        debug_log(f"[DEBUG][preview] /api/preview/households filters household_id={household_id!r} search={search!r} gender={gender!r} status={status!r} month={month!r}")
        cursor.execute(query, params)
        households = cursor.fetchall()

        items = []
        for h in households:
            items.append({
                'id': h.get('id'),
                'surname': h.get('surname'),
                'house_number': h.get('house_number'),
                'member_count': h.get('member_count', 0),
                'active_count': h.get('active_count', 0),
                'deceased_count': h.get('deceased_count', 0),
                'registration_date': str(h.get('registration_date')) if h.get('registration_date') else None
            })

        return jsonify({
            'total': len(items),
            'items': items
        })
    except Exception as e:
        print(f"[ERROR][preview] /api/preview/households error: {e}")
        return jsonify({'error': str(e), 'total': 0, 'items': []}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/preview/audit')
@require_role('admin')
def api_preview_audit():
    """Preview API for audit log reports"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    conn = None
    try:
        action_type = request.args.get('action_type', '').strip()
        target_type = request.args.get('target_type', '').strip()
        username = request.args.get('username', '').strip()
        month = request.args.get('month', '').strip()

        try:
            month_int = int(month) if month else None
            if month == 'year':
                month_int = 'year'
            elif month_int and (month_int < 1 or month_int > 12):
                month_int = None
        except ValueError:
            month_int = None

        ensure_audit_log_schema()
        conn, cursor = get_db()
        query = "SELECT * FROM audit_logs WHERE 1=1"
        params = []

        if action_type:
            query += " AND action_type = %s"
            params.append(action_type)
        if target_type:
            query += " AND target_type = %s"
            params.append(target_type)
        if username:
            query += " AND username = %s"
            params.append(username)
        if month_int == 'year':
            query += f" AND {_current_year_filter_sql('created_at')}"
        elif month_int:
            query += f" AND {_month_filter_sql('created_at')}"
            params.append(month_int)

        query += " ORDER BY created_at DESC LIMIT 100"
        debug_log(f"[DEBUG][preview] /api/preview/audit filters action_type={action_type!r} target_type={target_type!r} username={username!r} month={month!r}")
        cursor.execute(query, params)
        logs = cursor.fetchall()

        items = []
        for log in logs:
            items.append({
                'id': log.get('id'),
                'timestamp': str(log.get('created_at')) if log.get('created_at') else None,
                'action_type': log.get('action_type'),
                'target_type': log.get('target_type'),
                'description': log.get('details'),
                'username': log.get('username')
            })

        return jsonify({
            'total': len(items),
            'items': items
        })
    except Exception as e:
        print(f"[ERROR][preview] /api/preview/audit error: {e}")
        return jsonify({'error': str(e), 'total': 0, 'items': []}), 500
    finally:
        if conn:
            conn.close()

@app.route('/analytics')
@require_role('admin', 'user')
def analytics():
    if 'username' not in session:
        return redirect(url_for('login'))
    role = session.get('role', 'user')
    template = 'user_analytics.html' if role == 'user' else 'analytics.html'
    return render_template(template)

@app.route('/api/dashboard')
@require_role('admin', 'user')
def api_dashboard():
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    conn = None
    try:
        activity = request.args.get('activity', 'registered').strip()
        gender = request.args.get('gender', '').strip()
        status = request.args.get('status', '').strip()
        month = request.args.get('month', '').strip()
        date_from = request.args.get('date_from', '').strip()
        date_to = request.args.get('date_to', '').strip()
        household = request.args.get('household', '').strip()
        household_month = request.args.get('household_month', '').strip()
        household_from = request.args.get('household_from', '').strip()
        household_to = request.args.get('household_to', '').strip()

        if activity not in {'registered', 'deleted'}:
            activity = 'registered'

        def parse_filter_date(value):
            try:
                return datetime.strptime(value, '%Y-%m-%d').date() if value else None
            except ValueError:
                return None

        def iso_date(value):
            if not value:
                return None
            return value.isoformat() if hasattr(value, 'isoformat') else str(value)

        filter_from = parse_filter_date(date_from)
        filter_to = parse_filter_date(date_to)
        hh_filter_from = parse_filter_date(household_from)
        hh_filter_to = parse_filter_date(household_to)
        
        print('Reached /api/dashboard start')
        ensure_audit_log_schema()
        conn, cursor = get_db()
        print('Reached database connection')
        debug_log(f"\n[DEBUG][statistics] ===== /api/dashboard START =====")
        debug_log(f"[DEBUG][statistics] Filters: activity={activity!r} gender={gender!r} status={status!r} month={month!r}")
        debug_log(f"[DEBUG][statistics]   date_from={date_from!r} date_to={date_to!r}")
        debug_log(f"[DEBUG][statistics]   household={household!r} household_month={household_month!r}")
        debug_log(f"[DEBUG][statistics]   household_from={household_from!r} household_to={household_to!r}")
        
        if activity == 'deleted':
            print("Reached deleted resident statistics")
            cursor.execute("SELECT id, surname FROM households")
            household_names = {str(row['id']): row['surname'] for row in cursor.fetchall()}
            debug_log(f"[DEBUG][statistics] Household names loaded: {len(household_names)} households")
            
            deleted_query = """
                SELECT *
                FROM audit_logs
                WHERE target_type = 'Resident'
                  AND action_type = 'DELETE'
            """
            params = []
            if month:
                deleted_query += f" AND {_month_filter_sql('created_at')}"
                params.append(month)
            if filter_from:
                deleted_query += " AND DATE(created_at) >= %s"
                params.append(filter_from)
            if filter_to:
                deleted_query += " AND DATE(created_at) <= %s"
                params.append(filter_to)
            deleted_query += " ORDER BY created_at DESC"
            debug_log(f"[DEBUG][statistics] Deleted query: {deleted_query!r} with params {params!r}")
            cursor.execute(deleted_query, params)
            deleted_logs = cursor.fetchall()
            debug_log(f"[DEBUG][statistics] Deleted logs fetched: {len(deleted_logs)} records")
            residents = []
            for log in deleted_logs:
                snapshot = {}
                try:
                    snapshot = json.loads(log.get('old_value') or '{}')
                except Exception:
                    snapshot = {}
                if gender and snapshot.get('gender') != gender:
                    continue
                if status and (snapshot.get('status') or 'Active') != status:
                    continue
                hh_name = household_names.get(str(snapshot.get('household_id')), 'Unknown')
                if household and hh_name != household:
                    continue
                item = dict(snapshot)
                item['id'] = snapshot.get('id') or log.get('target_id')
                item['household_name'] = hh_name
                item['address'] = None
                item['registration_date'] = log.get('created_at')
                item['date_of_death'] = None
                item['status'] = snapshot.get('status') or 'Deleted'
                item['deleted_at'] = log.get('created_at')
                residents.append(item)
        else:
            print("Reached registration statistics")
            # Build resident query with real registration/deceased timestamps from audit logs.
            base_query = """
                SELECT *
                FROM (
                    SELECT
                        h.*,
                        hh.surname AS household_name,
                        hh.address AS address,
                        COALESCE((
                            SELECT MIN(al.created_at)
                            FROM audit_logs al
                            WHERE al.target_type = 'Resident'
                              AND al.target_id = CAST(h.id AS TEXT)
                              AND al.action_type = 'ADD'
                        ), h.created_at) AS registration_date,
                        (
                            SELECT MIN(al.created_at)
                            FROM audit_logs al
                            WHERE al.target_type = 'Resident'
                              AND al.target_id = CAST(h.id AS TEXT)
                              AND al.action_type = 'UPDATE'
                              AND (
                                  al.new_value ILIKE '%%"status": "Deceased"%%'
                                  OR al.new_value ILIKE '%%"status":"Deceased"%%'
                              )
                        ) AS date_of_death
                    FROM household h
                    LEFT JOIN households hh ON h.household_id = hh.id
                ) resident_dashboard
                WHERE 1=1
            """
            params = []
            
            if gender:
                base_query += " AND gender = %s"
                params.append(gender)
            
            if status:
                base_query += " AND COALESCE(status, 'Active') = %s"
                params.append(status)

            if household:
                base_query += " AND household_name = %s"
                params.append(household)
            
            if month:
                base_query += f" AND {_month_filter_sql('registration_date')}"
                params.append(month)

            if filter_from:
                base_query += " AND DATE(registration_date) >= %s"
                params.append(filter_from)

            if filter_to:
                base_query += " AND DATE(registration_date) <= %s"
                params.append(filter_to)
            
            debug_log(f"[DEBUG][statistics] Residents query SQL:\n{base_query}")
            debug_log(f"[DEBUG][statistics] Residents query params: {params!r}")
            print("SQL:", base_query)
            print("PARAMS:", params)
            print("PARAM COUNT:", len(params))
            print("PLACEHOLDER COUNT:", base_query.count("%s"))
            if len(params) != base_query.count("%s"):
                raise Exception(f"Resident query parameter mismatch: {len(params)} params for {base_query.count('%s')} placeholders")
            cursor.execute(base_query, params)
            residents = cursor.fetchall()
            debug_log(f"[DEBUG][statistics] Residents fetched: {len(residents)} records")

        household_query = """
            SELECT *
            FROM (
                SELECT
                    hh.id,
                    hh.surname AS household_name,
                    hh.house_number,
                    hh.address,
                    COUNT(h.id) AS members,
                    SUM(CASE WHEN COALESCE(h.status, 'Active') != 'Deceased' THEN 1 ELSE 0 END) AS active,
                    SUM(CASE WHEN COALESCE(h.status, 'Active') = 'Deceased' THEN 1 ELSE 0 END) AS deceased,
                    SUM(CASE WHEN h.gender = 'Male' THEN 1 ELSE 0 END) AS male,
                    SUM(CASE WHEN h.gender = 'Female' THEN 1 ELSE 0 END) AS female,
                    COALESCE((
                        SELECT MIN(al.created_at)
                        FROM audit_logs al
                        WHERE al.target_type = 'Household'
                          AND al.target_id = CAST(hh.id AS TEXT)
                          AND al.action_type = 'ADD'
                    ), hh.created_at) AS registration_date
                FROM households hh
                LEFT JOIN household h ON hh.id = h.household_id
                GROUP BY hh.id, hh.surname, hh.house_number, hh.address, hh.created_at
            ) household_dashboard
            WHERE 1=1
        """
        household_params = []

        if household:
            household_query += " AND household_name = %s"
            household_params.append(household)
        if household_month:
            household_query += f" AND {_month_filter_sql('registration_date')}"
            household_params.append(household_month)
        if hh_filter_from:
            household_query += " AND DATE(registration_date) >= %s"
            household_params.append(hh_filter_from)
        if hh_filter_to:
            household_query += " AND DATE(registration_date) <= %s"
            household_params.append(hh_filter_to)

        household_query += " ORDER BY household_name ASC"
        print('Reached household statistics')
        debug_log(f"[DEBUG][statistics] Households query SQL:\n{household_query}")
        debug_log(f"[DEBUG][statistics] Households query params: {household_params!r}")
        print("HOUSEHOLD SQL:", household_query)
        print("HOUSEHOLD PARAMS:", household_params)
        print("HOUSEHOLD PARAM COUNT:", len(household_params))
        print("HOUSEHOLD PLACEHOLDER COUNT:", household_query.count("%s"))
        if len(household_params) != household_query.count("%s"):
            raise Exception(f"Household query parameter mismatch: {len(household_params)} params for {household_query.count('%s')} placeholders")
        cursor.execute(household_query, household_params)
        households = cursor.fetchall()
        debug_log(f"[DEBUG][statistics] Households fetched: {len(households)} records")
        
        # Calculate stats
        stats = {
            'total': len(residents),
            'male': sum(1 for r in residents if r.get('gender') == 'Male'),
            'female': sum(1 for r in residents if r.get('gender') == 'Female'),
            'deceased': sum(1 for r in residents if r.get('status') == 'Deceased')
        }
        debug_log(f"[DEBUG][statistics] Stats calculated: {stats}")
        
        # Gender distribution
        gender_data = {
            'labels': ['Male', 'Female'],
            'datasets': [{
                'data': [stats['male'], stats['female']],
                'backgroundColor': ['#3b82f6', '#ec4899']
            }]
        }
        
        # Status distribution
        active = stats['total'] - stats['deceased']
        status_data = {
            'labels': ['Active', 'Deceased'],
            'datasets': [{
                'data': [active, stats['deceased']],
                'backgroundColor': ['#10b981', '#ef4444']
            }]
        }
        
        # Monthly registration data - separate by stat type
        month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

        def safe_month(value):
            if not value:
                return None
            if hasattr(value, 'month'):
                return value.month
            parts = str(value).split('-')
            if len(parts) < 2:
                return None
            try:
                m = int(parts[1])
            except ValueError:
                return None
            return m if 1 <= m <= 12 else None
        
        # Initialize counters for each stat type
        monthly_total = {month_names[i-1]: 0 for i in range(1, 13)}
        monthly_male = {month_names[i-1]: 0 for i in range(1, 13)}
        monthly_female = {month_names[i-1]: 0 for i in range(1, 13)}
        monthly_deceased = {month_names[i-1]: 0 for i in range(1, 13)}
        
        print("Reached deceased statistics")
        for r in residents:
            reg_date = r.get('registration_date')
            if reg_date:
                try:
                    m = safe_month(reg_date)
                    if m:
                        month_label = month_names[m-1]
                        monthly_total[month_label] += 1
                        if r.get('gender') == 'Male':
                            monthly_male[month_label] += 1
                        elif r.get('gender') == 'Female':
                            monthly_female[month_label] += 1
                except Exception:
                    pass

            death_date = r.get('date_of_death')
            if r.get('status') == 'Deceased' and death_date:
                try:
                    m = safe_month(death_date)
                    if m:
                        month_label = month_names[m-1]
                        monthly_deceased[month_label] += 1
                except Exception:
                    pass
        
        month_data = {
            'labels': month_names,
            'datasets': [
                {
                    'stat': 'total',
                    'label': 'Total Residents',
                    'data': [monthly_total[m] for m in month_names],
                    'backgroundColor': '#2e86c1'
                },
                {
                    'stat': 'male',
                    'label': 'Male Residents',
                    'data': [monthly_male[m] for m in month_names],
                    'backgroundColor': '#1B3A57'
                },
                {
                    'stat': 'female',
                    'label': 'Female Residents',
                    'data': [monthly_female[m] for m in month_names],
                    'backgroundColor': '#c0392b'
                },
                {
                    'stat': 'deceased',
                    'label': 'Deceased',
                    'data': [monthly_deceased[m] for m in month_names],
                    'backgroundColor': '#f59e0b'
                }
            ]
        }
        
        print('Reached response creation')
        # Top households
        household_counts = {}
        for r in residents:
            hh_name = r.get('household_name') or 'Unknown'
            household_counts[hh_name] = household_counts.get(hh_name, 0) + 1
        
        top_households = sorted(household_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        household_data = {
            'labels': [h[0][:15] for h in top_households],
            'datasets': [{
                'data': [h[1] for h in top_households],
                'backgroundColor': '#8b5cf6'
            }]
        }
        
        # Format residents for the frontend
        residents_list = []
        for r in residents:
            residents_list.append({
                'id': r.get('id'),
                'firstname': r.get('firstname'),
                'surname': r.get('surname'),
                'middlename': r.get('middlename'),
                'age': r.get('age'),
                'gender': r.get('gender'),
                'status': r.get('status', 'Active'),
                'household': r.get('household_name'),
                'address': r.get('address'),
                'birthdate': r.get('birthdate'),
                'registration_date': iso_date(r.get('registration_date')),
                'date_of_death': iso_date(r.get('date_of_death')),
                'deleted_at': iso_date(r.get('deleted_at')),
                'civil_status': r.get('civil_status'),
                'occupation': r.get('occupation')
            })

        households_list = []
        for h in households:
            households_list.append({
                'id': h.get('id'),
                'household_name': h.get('household_name'),
                'house_number': h.get('house_number'),
                'address': h.get('address'),
                'members': h.get('members') or 0,
                'active': h.get('active') or 0,
                'deceased': h.get('deceased') or 0,
                'male': h.get('male') or 0,
                'female': h.get('female') or 0,
                'registration_date': iso_date(h.get('registration_date'))
            })
        
        debug_log(f"[DEBUG][statistics] Formatted residents_list: {len(residents_list)} items")
        debug_log(f"[DEBUG][statistics] Formatted households_list: {len(households_list)} items")
        debug_log(f"[DEBUG][statistics] Month data labels: {month_data.get('labels')}")
        debug_log(f"[DEBUG][statistics] Month data datasets count: {len(month_data.get('datasets', []))}")
        if month_data.get('datasets'):
            for ds in month_data['datasets']:
                debug_log(f"[DEBUG][statistics]   Dataset stat={ds.get('stat')} label={ds.get('label')} data={ds.get('data')}")
        
        response = {
            'stats': stats,
            'genderData': gender_data,
            'statusData': status_data,
            'monthData': month_data,
            'householdData': household_data,
            'activity': activity,
            'residents': residents_list,
            'households': households_list
        }
        debug_log(f"[DEBUG][statistics] Final response keys: {list(response.keys())}")
        debug_log(f"[DEBUG][statistics] ===== /api/dashboard END =====\n")
        
        return jsonify(response)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print("Dashboard exception:", repr(e))
        raise
    finally:
        if conn:
            conn.close()

def init_app():
    """Initialize app - create all necessary tables at startup."""
    with app.app_context():
        try:
            if ensure_auth_schema_safe():
                print("Database schema initialized successfully.")
            else:
                print("Database schema initialization failed. Check DATABASE_URL and PostgreSQL logs.")
        except Exception as e:
            print(f"Database schema initialization skipped: {e}")

# Render/Gunicorn imports this module instead of executing the __main__ block.
# Running initialization here ensures the admin user is created in production too.
init_app()

if __name__ == '__main__':
    app.run(debug=True)
